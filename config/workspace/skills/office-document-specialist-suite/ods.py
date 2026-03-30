#!/usr/bin/env python3
# ref: https://python-docx.readthedocs.io
from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_paragraph_spacing(style, before=0, after=6, line=1.15):
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def ensure_style(doc: Document, name: str, base: str | None = None, size: int = 11, bold=False, color=None):
    styles = doc.styles
    if name in [s.name for s in styles]:
        style = styles[name]
    else:
        style = styles.add_style(name, 1)
        if base:
            style.base_style = styles[base]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)
    set_paragraph_spacing(style)
    return style


def insert_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def apply_advanced_layout(doc: Document, landscape=False):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    if landscape:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    header = section.header.paragraphs[0]
    header.text = "Clari AI"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    insert_page_number(footer)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    set_paragraph_spacing(normal, before=0, after=6, line=1.15)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    set_paragraph_spacing(h1, before=12, after=8, line=1.2)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2F, 0x55, 0x8C)
    set_paragraph_spacing(h2, before=10, after=6, line=1.2)

    ensure_style(doc, "Quote Accent", base="Normal", size=11, color=(85, 85, 85))


def cmd_template_report(output: Path, title: str, author: str):
    doc = Document()
    apply_advanced_layout(doc)
    configure_styles(doc)

    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph(f"작성자: {author}\n작성일: {date.today().isoformat()}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_paragraph("목차", style="Heading 1")
    doc.add_paragraph("(Word에서 목차 업데이트: 우클릭 > 필드 업데이트)")

    doc.add_page_break()
    doc.add_paragraph("1. 요약", style="Heading 1")
    doc.add_paragraph("여기에 주요 내용을 입력하세요.")

    doc.add_paragraph("2. 본문", style="Heading 1")
    doc.add_paragraph("2단계 제목과 3단계 제목으로 구조화하세요.")

    quote = doc.add_paragraph("주요 인용문이나 핵심 메시지.", style="Quote Accent")
    quote.paragraph_format.left_indent = Cm(1)

    doc.add_paragraph("2.1 세부 내용", style="Heading 2")
    doc.add_paragraph("표, 그림, 기술적 내용을 여기에 추가하세요.")

    doc.save(output)
    print(f"Template aangemaakt: {output}")


def cmd_style_doc(input_file: Path, output_file: Path, landscape: bool):
    doc = Document(input_file)
    apply_advanced_layout(doc, landscape=landscape)
    configure_styles(doc)
    doc.save(output_file)
    print(f"Opmaak toegepast: {output_file}")


def build_parser():
    parser = argparse.ArgumentParser(description="Office Document Specialist Suite")
    sub = parser.add_subparsers(dest="command", required=True)

    rep = sub.add_parser("template-report", help="Word 보고서 템플릿 생성 (.docx)")
    rep.add_argument("--output", default="report.docx", type=Path)
    rep.add_argument("--title", default="보고서 제목")
    rep.add_argument("--author", default="작성자")

    sty = sub.add_parser("style-doc", help="기존 .docx 파일에 스타일 적용")
    sty.add_argument("input", type=Path)
    sty.add_argument("--output", default="styled-output.docx", type=Path)
    sty.add_argument("--landscape", action="store_true")

    return parser


def main():
    parser = build_parser()
    args = parser.parse_args()

    if args.command == "template-report":
        cmd_template_report(args.output, args.title, args.author)
    elif args.command == "style-doc":
        cmd_style_doc(args.input, args.output, args.landscape)


if __name__ == "__main__":
    main()
